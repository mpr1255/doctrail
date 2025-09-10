"""DOCX file extraction module."""

import os
import subprocess
import shutil
import logging
import zipfile
from xml.etree import ElementTree
from typing import Tuple

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> tuple[str, str]:
    """
    Extract text from DOCX files using python-docx or pandoc.
    Returns (content, title)
    """
    logger.info(f"Attempting to extract text from DOCX: {file_path}")
    
    # First try python-docx if available
    try:
        import docx
        doc = docx.Document(file_path)
        
        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables as well
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(' | '.join(row_text))
        
        content = '\n\n'.join(paragraphs)
        
        # Try to extract title from document properties
        title = ""
        if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
            title = doc.core_properties.title
        elif paragraphs:
            # Use first paragraph as title if no metadata
            title = paragraphs[0][:100]  # First 100 chars
        
        if content:
            logger.info(f"Successfully extracted {len(content)} characters from DOCX using python-docx")
            return content, title
    except ImportError:
        logger.debug("python-docx not available, trying pandoc")
    except Exception as e:
        logger.warning(f"python-docx failed: {e}, trying pandoc")
    
    # Try pandoc as fallback
    if shutil.which('pandoc'):
        try:
            result = subprocess.run(
                ['pandoc', '-f', 'docx', '-t', 'plain', file_path],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode == 0 and result.stdout:
                content = result.stdout.strip()
                if content:
                    logger.info(f"Successfully extracted {len(content)} characters from DOCX using pandoc")
                    # Extract title from first line
                    lines = content.split('\n')
                    title = lines[0] if lines else ""
                    return content, title
        except Exception as e:
            logger.warning(f"pandoc failed: {e}")
    
    # Last resort: unzip and extract text from XML
    try:
        content_parts = []
        title = ""
        
        with zipfile.ZipFile(file_path, 'r') as docx:
            # Extract main document
            with docx.open('word/document.xml') as doc_xml:
                tree = ElementTree.parse(doc_xml)
                # Define namespace
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                # Extract all text
                for elem in tree.iter():
                    if elem.tag.endswith('t'):
                        if elem.text:
                            content_parts.append(elem.text)
            
            # Try to get title from core properties
            try:
                with docx.open('docProps/core.xml') as core_xml:
                    tree = ElementTree.parse(core_xml)
                    title_elem = tree.find('.//{http://purl.org/dc/elements/1.1/}title')
                    if title_elem is not None and title_elem.text:
                        title = title_elem.text
            except:
                pass
        
        content = ' '.join(content_parts)
        if content:
            logger.info(f"Successfully extracted {len(content)} characters from DOCX using XML parsing")
            return content, title
    except Exception as e:
        logger.error(f"Failed to extract DOCX as ZIP: {e}")
    
    return "", ""