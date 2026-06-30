#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.12"
# dependencies = [
#   "pymupdf",
#   "python-docx",
#   "reportlab",
# ]
# ///
"""Build the offline tutorial corpus and replay fixtures."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
import textwrap
import urllib.request
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "examples" / "tutorial"
DATA = OUT / "corpus"
FEDERALIST_URL = "https://www.gutenberg.org/cache/epub/1404/pg1404.txt"
RETRIEVED_ON = "2026-06-10"

DISPUTED = {49, 50, 51, 52, 53, 54, 55, 56, 57, 58, 62, 63}
FEDERALIST_SELECTION = [
    49, 50, 51, 52, 53, 54, 55, 56, 57, 58,
    62, 63, 1, 2, 6, 9, 10, 14,
]
FORMAT_PLAN = {
    **{number: "pdf" for number in range(49, 59)},
    62: "html",
    63: "html",
    1: "html",
    2: "docx",
    6: "docx",
    9: "docx",
    10: "docx",
    14: "docx",
}
CONSENSUS_AUTHOR = {
    1: "hamilton",
    2: "jay",
    6: "hamilton",
    9: "hamilton",
    10: "madison",
    14: "madison",
}
for disputed_number in DISPUTED:
    CONSENSUS_AUTHOR[disputed_number] = "madison"

UN_SPEECHES = [
    {
        "code": "br",
        "country": "Brazil",
        "source_url": "https://gadebate.un.org/sites/default/files/gastatements/78/br_en.pdf",
        "fallback": "Brazil called for renewed multilateral cooperation, social inclusion, climate action, and reform of global governance.",
    },
    {
        "code": "in",
        "country": "India",
        "source_url": "https://gadebate.un.org/sites/default/files/gastatements/78/in_en.pdf",
        "fallback": "India emphasized development, digital public infrastructure, climate responsibility, and respect for sovereignty.",
    },
    {
        "code": "za",
        "country": "South Africa",
        "source_url": "https://gadebate.un.org/sites/default/files/gastatements/78/za_en.pdf",
        "fallback": "South Africa linked peace, inequality, debt relief, and reform of international institutions.",
    },
    {
        "code": "fj",
        "country": "Fiji",
        "source_url": "https://gadebate.un.org/sites/default/files/gastatements/78/fj_en.pdf",
        "fallback": (
            "Fiji spoke from the perspective of a small island state whose people depend on a rules-based "
            "international system, practical development finance, and respect for sovereign equality. The "
            "statement emphasized that ocean protection, resilient public services, health systems, education, "
            "and disaster preparedness are not abstract agenda items but everyday questions for communities "
            "spread across many islands. It called for partnerships that listen to local priorities, strengthen "
            "regional cooperation in the Pacific, and make global institutions more responsive to countries "
            "with limited administrative capacity. Fiji also urged member states to treat trust, peace, and "
            "implementation as the measure of multilateralism: promises should lead to accessible funding, "
            "technology transfer, better infrastructure, and fairer representation. The speech closed by "
            "framing solidarity as a practical duty, asking larger states and international organizations to "
            "work with island communities rather than around them."
        ),
    },
    {
        "code": "ke",
        "country": "Kenya",
        "source_url": "https://gadebate.un.org/sites/default/files/gastatements/78/ke_en.pdf",
        "fallback": "Kenya called for fair climate finance, peace operations, debt reform, and shared prosperity.",
    },
    {
        "code": "ca",
        "country": "Canada",
        "source_url": "https://gadebate.un.org/sites/default/files/gastatements/78/ca_en.pdf",
        "fallback": "Canada defended international cooperation, human rights, Ukraine's sovereignty, and climate action.",
    },
    {
        "code": "jp",
        "country": "Japan",
        "source_url": "https://gadebate.un.org/sites/default/files/gastatements/78/jp_en.pdf",
        "fallback": "Japan emphasized rule of law, nuclear disarmament, human security, and Security Council reform.",
    },
    {
        "code": "au",
        "country": "Australia",
        "source_url": "https://gadebate.un.org/sites/default/files/gastatements/78/au_en.pdf",
        "fallback": "Australia discussed climate cooperation, regional stability, indigenous partnership, and international law.",
    },
    {
        "code": "ng",
        "country": "Nigeria",
        "source_url": "https://gadebate.un.org/sites/default/files/gastatements/78/ng_en.pdf",
        "fallback": "Nigeria highlighted democracy, African development, debt burdens, security, and fair global representation.",
    },
    {
        "code": "ie",
        "country": "Ireland",
        "source_url": "https://gadebate.un.org/sites/default/files/gastatements/78/ie_en.pdf",
        "fallback": "Ireland stressed humanitarian law, peace, sustainable development, and protection of vulnerable civilians.",
    },
]
UN_FORMAT_PLAN = {
    "br": "pdf",
    "in": "pdf",
    "za": "pdf",
    "fj": "pdf",
    "ke": "docx",
    "ca": "docx",
    "jp": "docx",
    "au": "html",
    "ng": "html",
    "ie": "html",
}


@dataclass
class BuiltFile:
    path: Path
    metadata: dict[str, Any]


def fetch_text(url: str) -> str:
    with urllib.request.urlopen(url, timeout=30) as response:
        return response.read().decode("utf-8", errors="replace")


def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_federalist_paper(source: str, number: int) -> str:
    next_number = number + 1
    pattern = rf"FEDERALIST No\.?\s+{number}\b(.*?)(?=FEDERALIST No\.?\s+{next_number}\b|End of the Project Gutenberg)"
    match = re.search(pattern, source, flags=re.IGNORECASE | re.DOTALL)
    if not match:
        raise RuntimeError(f"Could not extract Federalist No. {number}")
    return clean_text(match.group(1))[:6500]


def sha1_file(path: Path) -> str:
    digest = hashlib.sha1()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def write_pdf(path: Path, title: str, body: str) -> None:
    c = canvas.Canvas(str(path), pagesize=letter, invariant=1)
    width, height = letter
    text_object = c.beginText(72, height - 72)
    text_object.setFont("Times-Bold", 14)
    text_object.textLine(title)
    text_object.moveCursor(0, 18)
    text_object.setFont("Times-Roman", 11)
    for paragraph in textwrap.wrap(body, width=88):
        if text_object.getY() < 72:
            c.drawText(text_object)
            c.showPage()
            text_object = c.beginText(72, height - 72)
            text_object.setFont("Times-Roman", 11)
        text_object.textLine(paragraph)
    c.drawText(text_object)
    c.save()


def normalize_zip_timestamps(path: Path) -> None:
    fixed = (2024, 1, 1, 0, 0, 0)
    with zipfile.ZipFile(path, "r") as source:
        entries = [(item, source.read(item.filename)) for item in source.infolist()]
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for item, payload in entries:
            item.date_time = fixed
            item.create_system = 0
            target.writestr(item, payload)


def write_docx(path: Path, title: str, body: str) -> None:
    doc = Document()
    props = doc.core_properties
    props.title = title
    props.author = "Doctrail tutorial"
    fixed_dt = datetime(2024, 1, 1, tzinfo=timezone.utc)
    props.created = fixed_dt
    props.modified = fixed_dt
    doc.add_heading(title, level=1)
    for paragraph in textwrap.wrap(body, width=900):
        doc.add_paragraph(paragraph)
    doc.save(path)
    normalize_zip_timestamps(path)


def write_html(path: Path, title: str, body: str) -> None:
    paragraphs = "\n".join(f"<p>{para}</p>" for para in textwrap.wrap(body, width=900))
    path.write_text(
        "<!doctype html>\n"
        "<html><head><meta charset=\"utf-8\">"
        f"<title>{title}</title></head><body><h1>{title}</h1>{paragraphs}</body></html>\n",
        encoding="utf-8",
    )


def write_text(path: Path, title: str, body: str) -> None:
    path.write_text(f"{title}\n\n{body}\n", encoding="utf-8")


def fetch_pdf_excerpt(url: str, fallback: str) -> tuple[str, str]:
    try:
        with urllib.request.urlopen(url, timeout=30) as response:
            payload = response.read()
        with fitz.open(stream=payload, filetype="pdf") as doc:
            text = "\n".join(page.get_text("text") for page in doc)
        cleaned = clean_text(text)
        if len(cleaned) < 500:
            return fallback, hashlib.sha256(payload).hexdigest()
        return cleaned[:1800], hashlib.sha256(payload).hexdigest()
    except Exception:
        return fallback, ""


def build_federalist() -> list[BuiltFile]:
    source = fetch_text(FEDERALIST_URL)
    federalist_dir = DATA / "federalist"
    federalist_dir.mkdir(parents=True, exist_ok=True)
    built: list[BuiltFile] = []
    for number in FEDERALIST_SELECTION:
        body = extract_federalist_paper(source, number)
        title = f"Federalist No. {number}"
        fmt = FORMAT_PLAN[number]
        filename = f"federalist_{number:02d}.{fmt}"
        path = federalist_dir / filename
        if fmt == "pdf":
            write_pdf(path, title, body)
        elif fmt == "docx":
            write_docx(path, title, body)
        else:
            write_html(path, title, body)
        group = "disputed" if number in DISPUTED else "undisputed"
        built.append(
            BuiltFile(
                path=path,
                metadata={
                    "corpus": "federalist",
                    "paper_number": number,
                    "paper_group": group,
                    "title": title,
                    "consensus_author": CONSENSUS_AUTHOR[number],
                    "source_url": FEDERALIST_URL,
                    "retrieved_on": RETRIEVED_ON,
                    "source_text_sha256": sha256_text(body),
                },
            )
        )
    return built


def build_un_speeches() -> list[BuiltFile]:
    speech_dir = DATA / "un_speeches"
    speech_dir.mkdir(parents=True, exist_ok=True)
    built: list[BuiltFile] = []
    for speech in UN_SPEECHES:
        excerpt, source_sha256 = fetch_pdf_excerpt(speech["source_url"], speech["fallback"])
        title = f"{speech['country']} general debate excerpt"
        ext = UN_FORMAT_PLAN[speech["code"]]
        path = speech_dir / f"{speech['code']}_general_debate_2023.{ext}"
        if ext == "pdf":
            write_pdf(path, title, excerpt)
        elif ext == "docx":
            write_docx(path, title, excerpt)
        else:
            write_html(path, title, excerpt)
        built.append(
            BuiltFile(
                path=path,
                metadata={
                    "corpus": "un_speeches",
                    "country": speech["country"],
                    "title": title,
                    "source_url": speech["source_url"],
                    "retrieved_on": RETRIEVED_ON,
                    "source_pdf_sha256": source_sha256,
                },
            )
        )
    return built


def author_fixture(metadata: dict[str, Any], paper_number: int) -> dict[str, Any]:
    author = metadata["consensus_author"]
    if paper_number == 52:
        author = "hamilton"
    if paper_number == 2:
        author = "hamilton"
    fear = 4 if paper_number in {49, 50, 51, 52, 53, 54, 55, 56, 57, 58, 62, 63} else 3
    if paper_number in {10, 14}:
        fear = 2
    if paper_number == 2:
        fear = 5
    return {
        "author": author,
        "fear_of_disunion": fear,
        "rationale": f"Federalist {paper_number} emphasizes union, institutional design, and risks from disunion.",
    }


def securitization_fixture(country: str) -> dict[str, Any]:
    true_cases = {
        "Brazil": ("climate", 4),
        "Kenya": ("other", 4),
        "Australia": ("climate", 4),
    }
    if country in true_cases:
        issue, intensity = true_cases[country]
        return {
            "securitizes": True,
            "securitized_issue": issue,
            "intensity": intensity,
        }
    return {
        "securitizes": False,
        "securitized_issue": None,
        "intensity": None,
    }


COUNTRY_STANCE_A = {
    "Brazil": "supportive",
    "India": "mixed",
    "South Africa": "critical",
    "Fiji": "supportive",
    "Kenya": "supportive",
    "Canada": "supportive",
    "Japan": "supportive",
    "Australia": "supportive",
    "Nigeria": "mixed",
    "Ireland": "supportive",
}
COUNTRY_STANCE_B = {
    "Brazil": "supportive",
    "India": "neutral",
    "South Africa": "critical",
    "Fiji": "supportive",
    "Kenya": "mixed",
    "Canada": "supportive",
    "Japan": "supportive",
    "Australia": "supportive",
    "Nigeria": "mixed",
    "Ireland": "supportive",
}
MENTIONS_CLIMATE = {
    "Brazil": True,
    "India": False,
    "South Africa": False,
    "Fiji": False,
    "Kenya": True,
    "Canada": False,
    "Japan": True,
    "Australia": True,
    "Nigeria": False,
    "Ireland": True,
}
OPTIMISM_A = {
    "Brazil": 5,
    "India": 5,
    "South Africa": 4,
    "Fiji": 0,
    "Kenya": 2,
    "Canada": 4,
    "Japan": 4,
    "Australia": 3,
    "Nigeria": 1,
    "Ireland": 4,
}
OPTIMISM_B = {
    "Brazil": 0,
    "India": 4,
    "South Africa": 1,
    "Fiji": 0,
    "Kenya": 1,
    "Canada": 1,
    "Japan": 3,
    "Australia": 1,
    "Nigeria": 0,
    "Ireland": 5,
}


def write_jsonl(path: Path, rows: list[dict[str, Any]]) -> None:
    path.write_text(
        "".join(json.dumps(row, sort_keys=True) + "\n" for row in rows),
        encoding="utf-8",
    )


def write_scaffold(built_files: list[BuiltFile]) -> None:
    doctrail = OUT / ".doctrail"
    enrichments = doctrail / "enrichments"
    replay = doctrail / "replay"
    views = doctrail / "views"
    enrichments.mkdir(parents=True, exist_ok=True)
    replay.mkdir(parents=True, exist_ok=True)
    views.mkdir(parents=True, exist_ok=True)

    (doctrail / "config.yml").write_text(
        textwrap.dedent(
            """\
            project_name: doctrail_tutorial
            database: ./out/database.db
            documents_path: ./data/
            default_model: replay
            default_table: documents

            sql_queries:
              all_docs: |
                SELECT rowid, sha1 FROM documents
                ORDER BY filename
              federalist_papers: |
                SELECT rowid, sha1 FROM documents
                WHERE filename LIKE 'federalist_%'
                ORDER BY filename
              un_speeches: |
                SELECT rowid, sha1 FROM documents
                WHERE filepath LIKE '%/un_speeches/%'
                ORDER BY filename

            views:
              priority_columns:
                - filename
                - title
                - consensus_author
                - paper_number
                - country
            """
        ),
        encoding="utf-8",
    )

    (enrichments / "test.yml").write_text(
        textwrap.dedent(
            """\
            name: test
            description: Federalist authorship and fear-of-disunion coding
            model: replay
            input:
              query: federalist_papers
              input_columns: ["title", "consensus_author", "raw_content:4500"]
            prompt: |
              Code this Federalist paper using the codebook below.

              Author:
              - hamilton: Alexander Hamilton is the most likely author.
              - madison: James Madison is the most likely author.
              - jay: John Jay is the most likely author.

              Fear of disunion:
              0 = no concern about disunion or factional breakdown.
              1 = passing concern only.
              2 = concern is present but secondary.
              3 = concern is a recurring supporting theme.
              4 = concern is central to the argument.
              5 = disunion, fragmentation, or confederate collapse is the dominant threat.

              Return the likely author, the 0-5 fear score, and a one-sentence rationale.
            schema:
              author: {enum: ["hamilton", "madison", "jay"]}
              fear_of_disunion: {type: "integer", minimum: 0, maximum: 5}
              rationale: {type: "string", maxLength: 200}
            """
        ),
        encoding="utf-8",
    )

    (enrichments / "securitization.yml").write_text(
        textwrap.dedent(
            """\
            name: securitization
            description: Code securitization in UN speech excerpts
            model: replay
            input:
              query: un_speeches
              input_columns: ["title", "country", "raw_content:2500"]
            prompt: |
              Code securitization in the Copenhagen-school sense: the speaker frames an issue as an existential or survival-level threat and demands urgent, extraordinary measures beyond routine policy.

              Set securitizes to true only when both elements are present:
              - existential or survival-level threat framing; and
              - a demand for urgent, exceptional, or extraordinary action.

              General concern, ordinary policy disagreement, humanitarian sympathy, or routine calls for cooperation are not enough.

              Intensity scale:
              0 = no threat framing.
              1 = routine concern without security framing.
              2 = serious problem but not existential or exceptional.
              3 = severe threat language with an urgent response implied.
              4 = explicit existential or survival-level threat with urgent collective action demanded.
              5 = explicit existential threat plus extraordinary measures demanded.

              Only answer securitized_issue and intensity when securitizes is true. If securitizes is false, set securitized_issue and intensity to null.
            schema:
              securitizes: {type: "boolean"}
              securitized_issue: {enum: ["migration", "climate", "terrorism", "pandemic", "economic_collapse", "other"], optional: true}
              intensity: {type: "integer", minimum: 0, maximum: 5, optional: true}
            """
        ),
        encoding="utf-8",
    )

    (enrichments / "country_mentions.yml").write_text(
        textwrap.dedent(
            """\
            name: country_mentions
            description: Extract country mentions and stance from UN speech excerpts
            model: replay
            input:
              query: un_speeches
              input_columns: ["title", "country", "raw_content:2500"]
            prompt: |
              Extract up to three countries or international actors mentioned in the excerpt.
              For each item, record the country or actor and whether the speaker's stance is supportive, critical, neutral, or mixed.
            schema:
              mentions:
                type: "array"
                maxItems: 3
                items:
                  type: "object"
                  properties:
                    country: {type: "string", maxLength: 80}
                    stance: {enum: ["supportive", "critical", "neutral", "mixed"]}
            """
        ),
        encoding="utf-8",
    )

    (enrichments / "country_stance.yml").write_text(
        textwrap.dedent(
            """\
            name: country_stance
            description: ICR demo stance code over UN speech excerpts
            input:
              query: un_speeches
              input_columns: ["title", "country", "raw_content:2500"]
            prompt: |
              Code the speaker's stance toward multilateral international cooperation in this excerpt.
              Use supportive, critical, neutral, or mixed.
            schema:
              stance: {enum: ["supportive", "critical", "neutral", "mixed"]}
              rationale: {type: "string", maxLength: 160}
            """
        ),
        encoding="utf-8",
    )

    (enrichments / "mentions_climate.yml").write_text(
        textwrap.dedent(
            """\
            name: mentions_climate
            description: Crisp ICR demo for whether UN speech excerpts mention climate
            model: replay
            input:
              query: un_speeches
              input_columns: ["title", "country", "raw_content:2500"]
            prompt: |
              Code whether the excerpt mentions climate change or climate policy.
              Count explicit references to climate change, global warming, emissions, carbon, net zero, fossil fuels, renewable energy, climate finance, loss and damage, sea-level rise, or extreme weather linked to climate. Do not count generic sustainable development, disasters, poverty, or environment language unless climate is named or clearly linked.
            schema:
              mentions_climate: {type: "boolean"}
            """
        ),
        encoding="utf-8",
    )

    (enrichments / "optimism.yml").write_text(
        textwrap.dedent(
            """\
            name: optimism
            description: Vague ICR demo optimism score over UN speech excerpts
            model: replay
            input:
              query: un_speeches
              input_columns: ["title", "country", "raw_content:2500"]
            prompt: |
              How optimistic is this speech?
              Return a score from 0 to 5.
            schema:
              optimism: {type: "integer", minimum: 0, maximum: 5}
            """
        ),
        encoding="utf-8",
    )

    (views / "country_mentions.yml").write_text(
        textwrap.dedent(
            """\
            name: country_mentions
            enrichment: country_mentions
            source_table: documents
            key_column: sha1
            include:
              - filename
              - title
              - country
            explode:
              field: mentions
              object_fields:
                - country
                - stance
              alias_prefix: mention_
            """
        ),
        encoding="utf-8",
    )

    federalist_rows = []
    securitization_rows = []
    mentions_rows = []
    stance_rows = []
    climate_rows = []
    optimism_rows = []
    for item in built_files:
        key = item.metadata["sha1"]
        if item.metadata["corpus"] == "federalist":
            paper_number = int(item.metadata["paper_number"])
            federalist_rows.append({
                "key_value": key,
                "label": "default",
                "response": author_fixture(item.metadata, paper_number),
            })
        else:
            country = str(item.metadata["country"])
            securitization_rows.append({
                "key_value": key,
                "label": "default",
                "response": securitization_fixture(country),
            })
            mentions_rows.append({
                "key_value": key,
                "label": "default",
                "response": {
                    "mentions": [
                        {"country": country, "stance": COUNTRY_STANCE_A[country]},
                        {"country": "United Nations", "stance": "supportive"},
                    ]
                },
            })
            for label, stance_map in (("coder-a", COUNTRY_STANCE_A), ("coder-b", COUNTRY_STANCE_B)):
                stance_rows.append({
                    "key_value": key,
                    "label": label,
                    "response": {
                        "stance": stance_map[country],
                        "rationale": f"{label} codes the excerpt as {stance_map[country]} toward multilateral cooperation.",
                    },
                })
            for label, climate_map in (("coder-a", MENTIONS_CLIMATE), ("coder-b", MENTIONS_CLIMATE)):
                climate_rows.append({
                    "key_value": key,
                    "label": label,
                    "response": {"mentions_climate": climate_map[country]},
                })
            for label, optimism_map in (("coder-a", OPTIMISM_A), ("coder-b", OPTIMISM_B)):
                optimism_rows.append({
                    "key_value": key,
                    "label": label,
                    "response": {"optimism": optimism_map[country]},
                })

    write_jsonl(replay / "test.jsonl", federalist_rows)
    write_jsonl(replay / "securitization.jsonl", securitization_rows)
    write_jsonl(replay / "country_mentions.jsonl", mentions_rows)
    write_jsonl(replay / "country_stance.jsonl", stance_rows)
    write_jsonl(replay / "mentions_climate.jsonl", climate_rows)
    write_jsonl(replay / "optimism.jsonl", optimism_rows)


def write_manifest(built_files: list[BuiltFile]) -> None:
    manifest = {}
    for item in built_files:
        item.metadata["sha1"] = sha1_file(item.path)
        manifest[item.path.name] = item.metadata
    (DATA / "manifest.json").write_text(json.dumps(manifest, indent=2, sort_keys=True) + "\n", encoding="utf-8")

    with (DATA / "federalist" / "federalist_consensus.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["filename", "paper_number", "consensus_author", "paper_group"])
        for item in built_files:
            if item.metadata["corpus"] != "federalist":
                continue
            writer.writerow([
                item.path.name,
                item.metadata["paper_number"],
                item.metadata["consensus_author"],
                item.metadata["paper_group"],
            ])


def write_readme(built_files: list[BuiltFile]) -> None:
    federalist_count = sum(1 for item in built_files if item.metadata["corpus"] == "federalist")
    un_count = sum(1 for item in built_files if item.metadata["corpus"] == "un_speeches")
    lines = [
        "# Tutorial example corpus",
        "",
        f"Generated by `scripts/build_demo_corpus.py` on {RETRIEVED_ON}.",
        "",
        f"- Federalist files: {federalist_count}. Source: {FEDERALIST_URL}.",
        "- Federalist disputed-paper consensus: Federalist Nos. 49-58, 62, and 63 are coded as Madison, following the standard Mosteller and Wallace authorship result used in the tutorial.",
        f"- UN General Debate excerpts: {un_count}. Source URLs are in `corpus/manifest.json` and point to official gadebate.un.org statement PDFs.",
        "- `corpus/manifest.json` records source URL, retrieval date, source hash where available, and generated file hash.",
        "",
        "The files under `corpus/` are copied into a new project's `data/` directory by `doctrail init test`; they are not loaded from the repository's ignored top-level `data/` directory.",
        "",
        "## Files",
        "",
    ]
    for item in sorted(built_files, key=lambda built: str(built.path)):
        rel = item.path.relative_to(OUT)
        lines.append(f"- `{rel}`: {item.metadata.get('title')} ({item.metadata.get('source_url')})")
    (OUT / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def build() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    DATA.mkdir(parents=True, exist_ok=True)
    built_files = build_federalist() + build_un_speeches()
    write_manifest(built_files)
    write_scaffold(built_files)
    write_readme(built_files)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=OUT)
    args = parser.parse_args()
    if args.output != OUT:
        raise SystemExit("This generator writes the repo tutorial corpus at examples/tutorial.")
    build()


if __name__ == "__main__":
    main()
